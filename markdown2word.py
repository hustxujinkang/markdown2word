#!/usr/bin/env python
# -*- coding: utf-8 -*-

import argparse
import os
import markdown
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup
import re

def convert_markdown_to_word(markdown_file, output_file):
    """将Markdown文件转换为Word文档
    
    Args:
        markdown_file (str): Markdown文件的路径
        output_file (str): 输出Word文件的路径
    """
    # 检查输入文件是否存在
    if not os.path.exists(markdown_file):
        print(f"错误: 找不到文件 '{markdown_file}'")
        return False
    
    # 读取Markdown文件内容
    with open(markdown_file, 'r', encoding='utf-8') as f:
        markdown_text = f.read()
    
    # 将Markdown转换为HTML
    html = markdown.markdown(markdown_text, extensions=['tables', 'fenced_code', 'nl2br'])
    
    # 使用BeautifulSoup解析HTML
    soup = BeautifulSoup(html, 'html.parser')
    
    # 创建一个新的Word文档
    doc = Document()
    
    # 处理HTML元素并添加到Word文档
    process_html_elements(soup, doc)
    
    # 保存Word文档
    doc.save(output_file)
    print(f"成功将 '{markdown_file}' 转换为 '{output_file}'")
    return True

def process_html_elements(soup, doc):
    """处理HTML元素并添加到Word文档
    
    Args:
        soup (BeautifulSoup): BeautifulSoup解析的HTML
        doc (Document): Word文档对象
    """
    # 处理文档中的所有顶级元素
    for element in soup.children:
        if element.name is None:  # 纯文本
            if element.strip():
                p = doc.add_paragraph(element.strip())
        elif element.name == 'h1':
            p = doc.add_heading(element.text.strip(), level=1)
        elif element.name == 'h2':
            p = doc.add_heading(element.text.strip(), level=2)
        elif element.name == 'h3':
            p = doc.add_heading(element.text.strip(), level=3)
        elif element.name == 'h4':
            p = doc.add_heading(element.text.strip(), level=4)
        elif element.name == 'h5':
            p = doc.add_heading(element.text.strip(), level=5)
        elif element.name == 'h6':
            p = doc.add_heading(element.text.strip(), level=6)
        elif element.name == 'p':
            p = doc.add_paragraph()
            process_inline_elements(element, p)
        elif element.name == 'ul':
            process_list(element, doc, is_numbered=False)
        elif element.name == 'ol':
            process_list(element, doc, is_numbered=True)
        elif element.name == 'pre':
            code = element.find('code')
            if code:
                p = doc.add_paragraph()
                code_text = code.text.strip()
                p.add_run(code_text).font.name = 'Courier New'
                p.style = 'Code'
        elif element.name == 'blockquote':
            p = doc.add_paragraph()
            p.style = 'Quote'
            process_inline_elements(element, p)
        elif element.name == 'table':
            process_table(element, doc)

def process_inline_elements(element, paragraph):
    """处理内联HTML元素并添加到段落
    
    Args:
        element (Tag): HTML元素
        paragraph (Paragraph): Word段落对象
    """
    for child in element.children:
        if child.name is None:  # 纯文本
            if child.strip():
                paragraph.add_run(child.strip())
        elif child.name == 'strong' or child.name == 'b':
            run = paragraph.add_run(child.text.strip())
            run.bold = True
        elif child.name == 'em' or child.name == 'i':
            run = paragraph.add_run(child.text.strip())
            run.italic = True
        elif child.name == 'code':
            run = paragraph.add_run(child.text.strip())
            run.font.name = 'Courier New'
        elif child.name == 'a':
            run = paragraph.add_run(child.text.strip())
            run.underline = True
            run.font.color.rgb = RGBColor(0, 0, 255)
            # 添加超链接
            # 注意: python-docx目前不直接支持添加超链接，这里只是改变了文本样式
        elif child.name == 'br':
            paragraph.add_run('\n')
        elif child.name == 'img':
            # 注意: 这里简化处理，只添加图片描述
            alt_text = child.get('alt', '图片')
            paragraph.add_run(f"[{alt_text}]")

def process_list(list_element, doc, is_numbered=False, level=0):
    """处理列表元素并添加到Word文档
    
    Args:
        list_element (Tag): 列表HTML元素
        doc (Document): Word文档对象
        is_numbered (bool): 是否为有序列表
        level (int): 列表嵌套级别
    """
    for item in list_element.find_all('li', recursive=False):
        # 创建列表项段落
        p = doc.add_paragraph()
        
        # 构建样式名称
        style_name = f"List {'Number' if is_numbered else 'Bullet'} {level+1}"
        
        # 检查样式是否存在，如果不存在则使用默认样式或创建新样式
        try:
            p.style = style_name
        except KeyError:
            # 尝试使用替代样式
            if is_numbered:
                try:
                    p.style = 'List Number'
                except KeyError:
                    # 如果替代样式也不存在，使用普通段落样式
                    pass
            else:
                try:
                    p.style = 'List Bullet'
                except KeyError:
                    # 如果替代样式也不存在，使用普通段落样式
                    pass
        
        # 处理列表项内容
        for child in item.children:
            if child.name is None:  # 纯文本
                if child.strip():
                    p.add_run(child.strip())
            elif child.name in ['ul', 'ol']:
                # 处理嵌套列表
                process_list(child, doc, child.name == 'ol', level+1)
            else:
                # 处理其他内联元素
                process_inline_elements(child, p)

def process_table(table_element, doc):
    """处理表格元素并添加到Word文档
    
    Args:
        table_element (Tag): 表格HTML元素
        doc (Document): Word文档对象
    """
    # 获取表格行和列数
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # 计算列数（使用第一行的单元格数量）
    first_row = rows[0]
    cols = len(first_row.find_all(['th', 'td']))
    
    # 创建Word表格
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'
    
    # 填充表格内容
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            if j < cols:  # 确保不超出列数
                # 设置单元格内容
                table_cell = table.cell(i, j)
                table_cell.text = cell.text.strip()
                
                # 如果是表头，设置为粗体
                if cell.name == 'th' or i == 0:
                    for paragraph in table_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

def main():
    # 创建命令行参数解析器
    parser = argparse.ArgumentParser(description='将Markdown文件转换为Word文档')
    parser.add_argument('input', help='输入的Markdown文件路径')
    parser.add_argument('-o', '--output', help='输出的Word文件路径')
    
    # 解析命令行参数
    args = parser.parse_args()
    
    # 如果未指定输出文件，则使用输入文件名（更改扩展名为.docx）
    if not args.output:
        base_name = os.path.splitext(args.input)[0]
        args.output = f"{base_name}.docx"
    
    # 执行转换
    convert_markdown_to_word(args.input, args.output)

if __name__ == "__main__":
    main()