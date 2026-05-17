"""AST → docx.

Walks the Document AST and emits into a python-docx `Document`. If you pass
a reference.docx, we load it and clear its body so only your template's
styles, theme, sections, headers/footers remain. If not, python-docx's
built-in default template is used (it already ships with Heading 1..9,
Normal, List Bullet, List Number, Quote, Intense Quote).

Two font-application modes:
- builtin template: full override (we own the template, we own the fonts)
- user template:   fill-missing only (respect what the user set, just
  backfill the CJK slot when they forgot it)

This is controlled by `respect_template_fonts`. By default it's True iff a
user template was supplied.
"""
from __future__ import annotations

import base64
import os
import socket
import tempfile
import urllib.error
import urllib.request
from dataclasses import dataclass
from typing import List, Optional

from docx import Document as DocxDocument
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt

from . import ast as A
from .fonts import DEFAULTS, FontSpec, apply_font, style_has_eastasia_font


# ------------------------- configuration ----------------------------------

@dataclass
class FontProfile:
    body: FontSpec = None
    heading: FontSpec = None
    code: FontSpec = None

    def __post_init__(self):
        self.body = self.body or DEFAULTS["body"]
        self.heading = self.heading or DEFAULTS["heading"]
        self.code = self.code or DEFAULTS["code"]


# Sanity caps for fetched/decoded images. URL images get a strict timeout;
# data URIs and HTTP fetches both get a size ceiling so a stray giant
# resource can't blow up memory.
_HTTP_TIMEOUT_SEC = 10
_IMAGE_SIZE_LIMIT_BYTES = 10 * 1024 * 1024  # 10 MiB


# ---------------- style name resolution with fallback --------------------

def _style_or_fallback(doc, preferred: list) -> Optional[str]:
    """Return the first style name from `preferred` that actually exists."""
    existing = {s.name for s in doc.styles}
    for name in preferred:
        if name in existing:
            return name
    return None


# ---------------------------- main renderer -------------------------------

_BUILTIN_TEMPLATE = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "templates", "default_zh.docx"
)


class Renderer:
    def __init__(
        self,
        template: Optional[str] = None,
        fonts: Optional[FontProfile] = None,
        respect_template_fonts: Optional[bool] = None,
    ):
        # When the user explicitly passes a template, default to respecting
        # whatever fonts that template defines. When we fall back to the
        # builtin template, we own the styling end-to-end, so override.
        if respect_template_fonts is None:
            respect_template_fonts = template is not None
        self._respect = respect_template_fonts

        resolved = template or (_BUILTIN_TEMPLATE if os.path.exists(_BUILTIN_TEMPLATE) else None)
        self.doc = DocxDocument(resolved) if resolved else DocxDocument()
        self.fonts = fonts or FontProfile()
        if resolved:
            self._clear_body()
        self._cache_styles()
        # Track temp files (URL downloads, decoded data URIs) so we can
        # clean them up after render() finishes. Otherwise a long-running
        # caller leaks files into /tmp.
        self._temp_files: List[str] = []

    # purge the loaded template's body so we start clean
    def _clear_body(self) -> None:
        body = self.doc.element.body
        sectPr = body.find(qn("w:sectPr"))
        for child in list(body):
            if child is not sectPr:
                body.remove(child)

    def _cache_styles(self) -> None:
        d = self.doc
        self.sty_heading = [
            _style_or_fallback(d, [f"Heading {i}", f"标题 {i}"]) for i in range(1, 7)
        ]
        self.sty_normal  = _style_or_fallback(d, ["Normal", "正文"])
        self.sty_bullet  = _style_or_fallback(d, ["List Bullet", "无序列表"])
        self.sty_number  = _style_or_fallback(d, ["List Number", "有序列表"])
        self.sty_quote   = _style_or_fallback(d, ["Quote", "Intense Quote", "引用"])
        self.sty_code    = _style_or_fallback(d, ["Code Block", "Code", "代码块", "HTML Preformatted"])
        self.sty_caption = _style_or_fallback(d, ["Caption", "图注"])

        # Pre-compute whether each style we care about already declares a
        # CJK font. Under respect-template mode, if eastAsia is set we
        # leave the style alone entirely; if not, we backfill the CJK slot
        # at the run level so Chinese doesn't fall back to Calibri.
        existing = {s.name for s in d.styles}
        candidates = {self.sty_normal, self.sty_quote, self.sty_caption}
        candidates.update(self.sty_heading)
        self._cjk_set = {}
        for name in candidates:
            if name and name in existing:
                self._cjk_set[name] = style_has_eastasia_font(d.styles[name])

    # --------- entry point ---------

    def render(self, document: A.Document, output_path: str) -> None:
        self._apply_core_properties(document.metadata)
        for block in document.blocks:
            self._render_block(block)
        self.doc.save(output_path)
        self._cleanup_temp_files()

    def _apply_core_properties(self, metadata: dict) -> None:
        """Map frontmatter to Word's core document properties.

        These show up under File → Info in Word. We don't insert title /
        author into the body — that's the template's job (cover page,
        header, etc.). The mapping is intentionally narrow; users can put
        whatever they like in frontmatter, we only pick the keys that have
        a clean docx equivalent.
        """
        if not metadata:
            return
        cp = self.doc.core_properties
        mapping = {
            "title":       "title",
            "subject":     "subject",
            "author":      "author",
            "keywords":    "keywords",
            "description": "comments",
            "category":    "category",
            "version":     "version",
        }
        for md_key, prop in mapping.items():
            val = metadata.get(md_key)
            if val is None:
                continue
            try:
                setattr(cp, prop, str(val))
            except (AttributeError, ValueError):
                # Some properties have type constraints; silently skip
                # values we can't coerce rather than break the whole
                # render over a metadata typo.
                pass

    def _cleanup_temp_files(self) -> None:
        for path in self._temp_files:
            try:
                os.unlink(path)
            except OSError:
                pass
        self._temp_files.clear()

    # --------- block dispatch ---------

    def _render_block(self, block) -> None:
        if isinstance(block, A.Heading):
            self._render_heading(block)
        elif isinstance(block, A.Paragraph):
            self._render_paragraph(block)
        elif isinstance(block, A.CodeBlock):
            self._render_code(block)
        elif isinstance(block, A.Quote):
            self._render_quote(block)
        elif isinstance(block, A.BulletList):
            self._render_list(block, numbered=False, level=0)
        elif isinstance(block, A.OrderedList):
            self._render_list(block, numbered=True, level=0)
        elif isinstance(block, A.Table):
            self._render_table(block)
        elif isinstance(block, A.Image):
            self._render_block_image(block)
        elif isinstance(block, A.ThematicBreak):
            self._render_hr()

    # --------- individual blocks ---------

    def _render_heading(self, h: A.Heading) -> None:
        style_name = self.sty_heading[h.level - 1] or self.sty_normal
        p = self.doc.add_paragraph(style=self.doc.styles[style_name]) if style_name else self.doc.add_paragraph()
        self._write_inlines(p, h.inlines, font=self.fonts.heading, style_name=style_name)

    def _render_paragraph(self, p_node: A.Paragraph) -> None:
        p = self.doc.add_paragraph(style=self.doc.styles[self.sty_normal]) if self.sty_normal else self.doc.add_paragraph()
        self._write_inlines(p, p_node.inlines, font=self.fonts.body, style_name=self.sty_normal)

    def _render_code(self, cb: A.CodeBlock) -> None:
        style = self.doc.styles[self.sty_code] if self.sty_code else (self.doc.styles[self.sty_normal] if self.sty_normal else None)
        p = self.doc.add_paragraph(style=style) if style else self.doc.add_paragraph()
        # preserve line breaks inside a single paragraph
        lines = cb.code.split("\n")
        for i, line in enumerate(lines):
            run = p.add_run(line)
            # Code is always full-override for the font (monospace must be
            # explicit — we can't trust the template to have a Code Block
            # style with a fixed-width font wired up).
            apply_font(run, self.fonts.code)
            if i < len(lines) - 1:
                run.add_break(WD_BREAK.LINE)

    def _render_quote(self, q: A.Quote) -> None:
        """Render a blockquote.

        The original implementation only restyled `doc.paragraphs[-1]`
        after each child, which silently dropped the Quote style on every
        non-last paragraph (e.g. multi-paragraph quotes, or quotes
        wrapping a list). Fix: snapshot the paragraph count before, then
        restyle every paragraph we appended whose current style is still
        Normal — that way list/code paragraphs inside a quote keep their
        own styles (which is what Word users expect) while plain
        paragraphs get tagged as Quote.
        """
        if not self.sty_quote:
            for child in q.children:
                self._render_block(child)
            return

        before = len(self.doc.paragraphs)
        for child in q.children:
            self._render_block(child)
        after = len(self.doc.paragraphs)

        quote_style = self.doc.styles[self.sty_quote]
        normal_name = self.sty_normal
        for p in self.doc.paragraphs[before:after]:
            current = p.style.name if p.style else ""
            if not normal_name or current == normal_name or current == "":
                p.style = quote_style

    def _render_list(self, node, numbered: bool, level: int) -> None:
        style_name = self.sty_number if numbered else self.sty_bullet
        for item in node.items:
            first = True
            for child in item.children:
                if isinstance(child, (A.BulletList, A.OrderedList)):
                    self._render_list(child, isinstance(child, A.OrderedList), level + 1)
                else:
                    if isinstance(child, A.Paragraph) and first:
                        p = self.doc.add_paragraph(
                            style=self.doc.styles[style_name]) if style_name else self.doc.add_paragraph()
                        self._write_inlines(
                            p, child.inlines,
                            font=self.fonts.body,
                            style_name=style_name,
                        )
                        if level > 0:
                            # Two things here:
                            # 1) Backfill an explicit left indent so deep
                            #    levels still look indented even when the
                            #    style's own numbering definition has no
                            #    per-level indent.
                            # 2) Inject w:numPr/w:ilvl so that *if* the
                            #    user's template defines a multi-level
                            #    list (1. / 1.1 / 1.1.1) under the same
                            #    numId as List Number, Word actually
                            #    renders the right marker for this level.
                            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
                            _set_list_level(p, level)
                        first = False
                    else:
                        self._render_block(child)

    def _render_table(self, t: A.Table) -> None:
        header_cells = t.header.cells if t.header else None
        ncols = len(header_cells) if header_cells else (len(t.rows[0].cells) if t.rows else 0)
        if ncols == 0:
            return
        total_rows = (1 if header_cells else 0) + len(t.rows)
        table = self.doc.add_table(rows=total_rows, cols=ncols)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        r_idx = 0
        if header_cells:
            for j, inlines in enumerate(header_cells):
                cell = table.cell(0, j)
                cell.text = ""
                p = cell.paragraphs[0]
                self._write_inlines(p, inlines, font=self.fonts.body,
                                    force_bold=True, style_name=self.sty_normal)
            r_idx = 1
        for row in t.rows:
            for j, inlines in enumerate(row.cells):
                if j >= ncols:
                    break
                cell = table.cell(r_idx, j)
                cell.text = ""
                p = cell.paragraphs[0]
                self._write_inlines(p, inlines, font=self.fonts.body,
                                    style_name=self.sty_normal)
            r_idx += 1

    def _render_block_image(self, img: A.Image) -> None:
        p = self.doc.add_paragraph()
        run = p.add_run()
        path = self._resolve_image(img.src)
        if path:
            try:
                run.add_picture(path, width=self._content_width())
            except Exception:
                run.add_text(f"[图片: {img.alt or img.src}]")
        else:
            run.add_text(f"[图片: {img.alt or img.src}]")
        if img.alt and self.sty_caption:
            cap = self.doc.add_paragraph(img.alt, style=self.doc.styles[self.sty_caption])
            for r in cap.runs:
                self._apply_run_font(r, self.fonts.body, style_name=self.sty_caption)

    def _render_hr(self) -> None:
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # --------- inlines ---------

    def _write_inlines(
        self,
        p,
        inlines,
        font: FontSpec,
        force_bold: bool = False,
        style_name: Optional[str] = None,
    ) -> None:
        for node in inlines:
            if isinstance(node, A.Text):
                if node.value == "\n":
                    p.add_run().add_break(WD_BREAK.LINE); continue
                run = p.add_run(node.value)
                run.bold = node.bold or force_bold
                run.italic = node.italic
                if node.code:
                    apply_font(run, self.fonts.code)
                else:
                    self._apply_run_font(run, font, style_name=style_name)
            elif isinstance(node, A.Link):
                self._add_hyperlink(p, node.url, node.text, font, style_name=style_name)
            elif isinstance(node, A.InlineImage):
                run = p.add_run()
                path = self._resolve_image(node.src)
                if path:
                    try:
                        run.add_picture(path, height=Pt(14))
                    except Exception:
                        run.add_text(f"[{node.alt or 'image'}]")
                else:
                    run.add_text(f"[{node.alt or 'image'}]")

    def _apply_run_font(self, run, font: FontSpec, style_name: Optional[str]) -> None:
        """Apply font to a run, honoring respect-template mode.

        - Builtin template path (respect=False): full override, as before.
        - User template path (respect=True): if the style already declares
          a CJK font, do nothing (trust the template). Otherwise backfill
          only the slots not already set on the run.
        """
        if not self._respect:
            apply_font(run, font)
            return
        if style_name and self._cjk_set.get(style_name):
            return
        apply_font(run, font, fill_missing_only=True)

    def _add_hyperlink(self, paragraph, url: str, text: str, font: FontSpec,
                       style_name: Optional[str] = None) -> None:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle"); rStyle.set(qn("w:val"), "Hyperlink"); rPr.append(rStyle)
        run.append(rPr)
        t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"), "preserve"); run.append(t)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

        # Backfill fonts on the hyperlink run. Under respect-template
        # mode, only do this when the surrounding paragraph style doesn't
        # already define a CJK font.
        if self._respect and style_name and self._cjk_set.get(style_name):
            return
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), font.latin); rfonts.set(qn("w:hAnsi"), font.latin)
        rfonts.set(qn("w:eastAsia"), font.cjk); rfonts.set(qn("w:cs"), font.latin)
        rPr.insert(0, rfonts)

    # --------- helpers ---------

    def _content_width(self) -> Emu:
        section = self.doc.sections[0]
        return Emu(section.page_width - section.left_margin - section.right_margin)

    # ----- image resolution (instance method so we can track temp files) -----

    def _resolve_image(self, src: str) -> Optional[str]:
        """Return a local path python-docx can embed, or None.

        Supports:
        - http(s)://: downloaded with a hard timeout and size cap. The
          downloaded file is tracked for cleanup after render().
        - data:image/...;base64,...: decoded to a temp file, tracked
          likewise. LLM-generated diagrams (Mermaid SSR, KaTeX, etc.)
          commonly arrive as data URIs, so this isn't an edge case.
        - any other path: returned if the file exists, else None.
        """
        if not src:
            return None

        if src.startswith("data:"):
            tmp = _decode_data_uri(src)
            if tmp:
                self._temp_files.append(tmp)
            return tmp

        if src.startswith(("http://", "https://")):
            tmp = _download_image(src)
            if tmp:
                self._temp_files.append(tmp)
            return tmp

        return src if os.path.exists(src) else None


# ---------------------- module-level helpers --------------------------

def _set_list_level(p, level: int) -> None:
    """Add or update a w:numPr/w:ilvl on the paragraph.

    We don't synthesize a numId here — the paragraph's style already
    carries one via its numbering definition. We just stamp the level so
    Word can pick the right per-level marker if the style's numbering
    defines one. If the template's list style is single-level, ilvl is
    harmless; if it's multi-level, this is what makes it work.
    """
    pPr = p._p.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)
    ilvl = numPr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        numPr.insert(0, ilvl)
    ilvl.set(qn("w:val"), str(level))


_DATA_URI_PREFIX = "data:"


def _decode_data_uri(src: str) -> Optional[str]:
    """Decode a data: URI image to a temp file. Returns path or None."""
    try:
        # data:[<mediatype>][;base64],<data>
        head, _, payload = src.partition(",")
        if not payload:
            return None
        meta = head[len(_DATA_URI_PREFIX):]
        parts = meta.split(";")
        media = parts[0] if parts else ""
        is_b64 = any(p == "base64" for p in parts[1:])
        if not is_b64:
            # We only support base64-encoded image data URIs. Percent-
            # encoded text payloads aren't useful as images.
            return None
        # Guess the extension from the media type so python-docx /
        # Word picks the right content-type.
        ext = ".png"
        if "/" in media:
            subtype = media.split("/", 1)[1].lower()
            if subtype in {"jpeg", "jpg"}:
                ext = ".jpg"
            elif subtype == "gif":
                ext = ".gif"
            elif subtype == "webp":
                ext = ".webp"
            elif subtype == "bmp":
                ext = ".bmp"
            elif subtype == "svg+xml":
                # python-docx can't embed SVG natively. Skip rather than
                # write a broken image into the doc.
                return None
        raw = base64.b64decode(payload, validate=False)
        if len(raw) > _IMAGE_SIZE_LIMIT_BYTES:
            return None
        fd, tmp = tempfile.mkstemp(suffix=ext)
        with os.fdopen(fd, "wb") as f:
            f.write(raw)
        return tmp
    except (ValueError, OSError, base64.binascii.Error):
        return None


def _download_image(url: str) -> Optional[str]:
    """Fetch an http(s) image with timeout + size cap. Returns path or None."""
    tmp = None
    try:
        suffix = os.path.splitext(url.split("?", 1)[0])[1] or ".png"
        fd, tmp = tempfile.mkstemp(suffix=suffix)
        os.close(fd)
        req = urllib.request.Request(url, headers={"User-Agent": "md2word/0.1"})
        with urllib.request.urlopen(req, timeout=_HTTP_TIMEOUT_SEC) as resp, open(tmp, "wb") as out:
            # Stream-copy with a size guard. Don't trust Content-Length —
            # some servers lie about it and some omit it entirely.
            total = 0
            while True:
                chunk = resp.read(64 * 1024)
                if not chunk:
                    break
                total += len(chunk)
                if total > _IMAGE_SIZE_LIMIT_BYTES:
                    out.close()
                    os.unlink(tmp)
                    return None
                out.write(chunk)
        return tmp
    except (urllib.error.URLError, socket.timeout, OSError, ValueError):
        if tmp is not None:
            try:
                os.unlink(tmp)
            except OSError:
                pass
        return None
