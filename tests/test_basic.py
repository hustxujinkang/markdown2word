"""Structural tests for md2word.

Each test is one feature. We assert on docx structure and XML, not on visual
appearance — visuals you verify by opening the file in Word.

Run with:
    python -m unittest discover tests

Each test:
  1. Reads tests/fixtures/<name>.md
  2. Pipes through parse() and Renderer into a temp .docx
  3. Reopens the .docx and asserts on paragraphs / runs / tables / XML
"""
from __future__ import annotations

import os
import tempfile
import unittest

from docx import Document
from docx.oxml.ns import qn

from md2word.parser import parse
from md2word.renderer import Renderer


FIXTURES = os.path.join(os.path.dirname(__file__), "fixtures")


# --------------------------- helpers --------------------------------------

def render_fixture(name: str):
    """Render tests/fixtures/<name>.md and return an opened Document."""
    md_path = os.path.join(FIXTURES, f"{name}.md")
    with open(md_path, encoding="utf-8") as f:
        md_text = f.read()
    ast = parse(md_text)
    fd, out = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    Renderer().render(ast, out)
    return Document(out), out


def style_of(paragraph) -> str:
    return paragraph.style.name if paragraph.style else ""


def run_fonts(run) -> dict:
    """Return {ascii, hAnsi, eastAsia, cs} for a run, or {} if not set."""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return {}
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return {}
    return {
        slot: rfonts.get(qn(f"w:{slot}"))
        for slot in ("ascii", "hAnsi", "eastAsia", "cs")
    }


def hyperlinks_in(paragraph) -> list:
    """Return [(url, text)] for every w:hyperlink in the paragraph."""
    out = []
    for h in paragraph._p.findall(qn("w:hyperlink")):
        r_id = h.get(qn("r:id"))
        target = paragraph.part.rels[r_id].target_ref if r_id else ""
        text = "".join(t.text or "" for t in h.iter(qn("w:t")))
        out.append((target, text))
    return out


def table_as_grid(table) -> list:
    return [[cell.text for cell in row.cells] for row in table.rows]


# --------------------------- tests ----------------------------------------

class TestHeadings(unittest.TestCase):
    def test_levels_1_through_6(self):
        doc, _ = render_fixture("headings")
        levels = [style_of(p) for p in doc.paragraphs if p.text]
        expected = [f"Heading {i}" for i in range(1, 7)]
        self.assertEqual(levels, expected)

    def test_heading_text_preserved(self):
        doc, _ = render_fixture("headings")
        texts = [p.text for p in doc.paragraphs if p.text]
        self.assertEqual(texts, ["一级", "二级", "三级", "四级", "五级", "六级"])


class TestInlineFormatting(unittest.TestCase):
    def test_bold_italic_code_runs(self):
        doc, _ = render_fixture("inline")
        p = next(p for p in doc.paragraphs if p.text)
        # concatenated runs should reconstruct the paragraph
        self.assertIn("加粗", p.text)
        self.assertIn("斜体", p.text)
        self.assertIn("行内代码", p.text)

        # at least one bold run and one italic run must exist
        has_bold = any(r.bold for r in p.runs)
        has_italic = any(r.italic for r in p.runs)
        self.assertTrue(has_bold, "expected at least one bold run")
        self.assertTrue(has_italic, "expected at least one italic run")

    def test_inline_code_uses_code_font(self):
        doc, _ = render_fixture("inline")
        p = next(p for p in doc.paragraphs if p.text)
        code_run = next((r for r in p.runs if "行内代码" in r.text), None)
        self.assertIsNotNone(code_run)
        fonts = run_fonts(code_run)
        self.assertEqual(fonts.get("ascii"), "Consolas")

    def test_combined_bold_italic(self):
        doc, _ = render_fixture("inline")
        p = next(p for p in doc.paragraphs if p.text)
        combined = next((r for r in p.runs if "又粗又斜" in r.text), None)
        self.assertIsNotNone(combined)
        self.assertTrue(combined.bold)
        self.assertTrue(combined.italic)


class TestLists(unittest.TestCase):
    def test_bullet_list_style(self):
        doc, _ = render_fixture("lists")
        bullet_texts = [p.text for p in doc.paragraphs if style_of(p) == "List Bullet"]
        self.assertEqual(bullet_texts, ["苹果", "香蕉", "橘子"])

    def test_ordered_list_style(self):
        doc, _ = render_fixture("lists")
        number_texts = [p.text for p in doc.paragraphs if style_of(p) == "List Number"]
        self.assertEqual(number_texts, ["第一条", "第二条", "第三条"])


class TestNestedList(unittest.TestCase):
    def test_all_items_emitted(self):
        doc, _ = render_fixture("nested_list")
        items = [p.text for p in doc.paragraphs if p.text]
        expected = ["第一层甲", "第二层甲一", "第二层甲二",
                    "第一层乙", "第二层乙一", "第三层"]
        self.assertEqual(items, expected)

    def test_deeper_levels_are_indented(self):
        doc, _ = render_fixture("nested_list")
        by_text = {p.text: p for p in doc.paragraphs if p.text}
        level1 = by_text["第一层甲"].paragraph_format.left_indent
        level3 = by_text["第三层"].paragraph_format.left_indent
        # level1 is 0 or None (uses style's own indent); level3 got explicit extra indent
        l1 = level1.emu if level1 else 0
        l3 = level3.emu if level3 else 0
        self.assertGreater(l3, l1, "deeper list levels should be indented further")


class TestTable(unittest.TestCase):
    def test_dimensions_and_content(self):
        doc, _ = render_fixture("table")
        self.assertEqual(len(doc.tables), 1)
        grid = table_as_grid(doc.tables[0])
        self.assertEqual(grid, [
            ["参数", "值", "说明"],
            ["字号", "12pt", "小四"],
            ["行距", "1.5", "倍数"],
        ])

    def test_header_is_bold(self):
        doc, _ = render_fixture("table")
        header_row = doc.tables[0].rows[0]
        for cell in header_row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if r.text.strip():
                        self.assertTrue(r.bold, f"header cell run {r.text!r} not bold")


class TestCodeBlock(unittest.TestCase):
    def test_uses_code_block_style(self):
        doc, _ = render_fixture("code")
        code_paras = [p for p in doc.paragraphs if style_of(p) == "Code Block"]
        self.assertEqual(len(code_paras), 1)

    def test_multiline_content_preserved(self):
        doc, _ = render_fixture("code")
        p = next(p for p in doc.paragraphs if style_of(p) == "Code Block")
        self.assertIn("def hello():", p.text)
        self.assertIn("你好", p.text)
        # should have line breaks inside (multiple w:br)
        breaks = p._p.findall(".//" + qn("w:br"))
        self.assertGreater(len(breaks), 0, "code block should preserve line breaks")

    def test_uses_monospace_font(self):
        doc, _ = render_fixture("code")
        p = next(p for p in doc.paragraphs if style_of(p) == "Code Block")
        r = p.runs[0]
        self.assertEqual(run_fonts(r).get("ascii"), "Consolas")


class TestQuote(unittest.TestCase):
    def test_uses_quote_style(self):
        doc, _ = render_fixture("quote")
        quote_paras = [p for p in doc.paragraphs if style_of(p) == "Quote"]
        self.assertGreater(len(quote_paras), 0)


class TestHyperlink(unittest.TestCase):
    def test_real_hyperlink_with_relationship(self):
        doc, _ = render_fixture("link")
        p = next(p for p in doc.paragraphs if p.text)
        links = hyperlinks_in(p)
        self.assertEqual(len(links), 1)
        url, text = links[0]
        self.assertEqual(url, "https://www.python.org/")
        self.assertEqual(text, "Python 官网")


class TestMissingImage(unittest.TestCase):
    def test_degrades_to_alt_text(self):
        doc, _ = render_fixture("missing_image")
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("[图片: 架构图]", full_text)
        # shouldn't have crashed — rest of doc is there too
        self.assertIn("降级为文字占位符", full_text)


class TestCJKFontSlots(unittest.TestCase):
    def test_both_latin_and_eastasia_slots_set(self):
        doc, _ = render_fixture("cjk_mixed")
        p = next(p for p in doc.paragraphs if p.text)
        self.assertGreater(len(p.runs), 0)
        for r in p.runs:
            fonts = run_fonts(r)
            self.assertEqual(fonts.get("ascii"), "Times New Roman",
                             f"run {r.text!r} missing Latin font")
            self.assertEqual(fonts.get("eastAsia"), "宋体",
                             f"run {r.text!r} missing CJK font")


class TestThematicBreak(unittest.TestCase):
    def test_hr_rendered_as_bottom_border(self):
        doc, _ = render_fixture("hr")
        # should have three paragraphs: "上面的段落" / HR / "下面的段落"
        found_border = False
        for p in doc.paragraphs:
            pPr = p._p.find(qn("w:pPr"))
            if pPr is None:
                continue
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is not None and pBdr.find(qn("w:bottom")) is not None:
                found_border = True
                break
        self.assertTrue(found_border, "expected at least one paragraph with bottom border")


class TestEndToEndSample(unittest.TestCase):
    """Smoke test: the comprehensive sample.md renders without errors."""

    def test_sample_renders(self):
        doc, out_path = render_fixture("sample")
        self.assertGreater(len(doc.paragraphs), 10)
        self.assertGreater(len(doc.tables), 0)
        os.unlink(out_path)


# ----------------- new tests: bug fixes & v0.2 features -----------------


class TestFrontmatter(unittest.TestCase):
    """YAML frontmatter is parsed out and routed to docx core properties."""

    def test_frontmatter_populates_core_properties(self):
        doc, out_path = render_fixture("frontmatter")
        cp = doc.core_properties
        self.assertEqual(cp.title, "二〇二六年第一季度工作总结")
        self.assertEqual(cp.author, "张三")
        self.assertEqual(cp.keywords, "季度总结, 工作回顾")
        self.assertEqual(cp.comments, "季度工作回顾与下阶段计划")
        os.unlink(out_path)

    def test_frontmatter_not_in_body(self):
        """Frontmatter shouldn't leak as text into the document body."""
        doc, out_path = render_fixture("frontmatter")
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertNotIn("title:", full_text)
        self.assertNotIn("author:", full_text)
        # The actual H1 from after the frontmatter should still be there.
        self.assertIn("第一季度总结", full_text)
        os.unlink(out_path)

    def test_no_frontmatter_is_harmless(self):
        """Documents without frontmatter still render normally."""
        from md2word.parser import parse as parse_md
        ast = parse_md("# 普通标题\n\n正文段落。\n")
        self.assertEqual(ast.metadata, {})
        self.assertEqual(len(ast.blocks), 2)

    def test_frontmatter_parser_directly(self):
        from md2word.parser import _parse_frontmatter
        md = "---\ntitle: 测试\nauthor: 李四\n---\n\n# 标题\n\n正文\n"
        meta, rest = _parse_frontmatter(md)
        self.assertEqual(meta["title"], "测试")
        self.assertEqual(meta["author"], "李四")
        self.assertTrue(rest.startswith("\n# 标题"))


class TestMultiParagraphQuote(unittest.TestCase):
    """Regression: every paragraph inside a blockquote must get Quote style,
    not just the last one."""

    def test_all_three_paragraphs_styled_as_quote(self):
        doc, out_path = render_fixture("multi_quote")
        quote_paras = [p for p in doc.paragraphs if style_of(p) == "Quote" and p.text]
        self.assertEqual(len(quote_paras), 3,
                         f"expected 3 Quote-styled paragraphs, got {len(quote_paras)}")
        texts = [p.text for p in quote_paras]
        self.assertIn("这是第一段引用。", texts)
        self.assertIn("这是第三段。", texts)
        os.unlink(out_path)


class TestDataUriImage(unittest.TestCase):
    """Inline base64 data URIs should be decoded and embedded as images,
    not dropped to alt text."""

    def test_data_uri_embedded(self):
        doc, out_path = render_fixture("data_uri_image")
        # The presence of an embedded image part means decode + embed worked
        rels = doc.part.rels
        image_rels = [r for r in rels.values() if "image" in r.reltype]
        self.assertEqual(len(image_rels), 1,
                         "expected one embedded image from the data URI")
        os.unlink(out_path)

    def test_decode_data_uri_unit(self):
        """Direct test of the decoder, independent of the renderer."""
        from md2word.renderer import _decode_data_uri
        # 1x1 transparent PNG
        png = ("data:image/png;base64,"
               "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNkYAAAAAYAAjCB0C8AAAAASUVORK5CYII=")
        tmp = _decode_data_uri(png)
        self.assertIsNotNone(tmp)
        self.assertTrue(os.path.exists(tmp))
        # First bytes should be the PNG magic
        with open(tmp, "rb") as f:
            head = f.read(4)
        self.assertEqual(head, b"\x89PNG")
        os.unlink(tmp)

    def test_decode_rejects_svg(self):
        """SVG data URIs are skipped (python-docx can't embed them)."""
        from md2word.renderer import _decode_data_uri
        svg = "data:image/svg+xml;base64,PHN2Zy8+"
        self.assertIsNone(_decode_data_uri(svg))


class TestDeepListLevels(unittest.TestCase):
    """Nested lists past level 0 should carry an explicit w:ilvl so that
    multi-level numbering in a user template can fire."""

    def test_ilvl_present_at_each_depth(self):
        doc, out_path = render_fixture("deep_list")
        levels_seen = {}
        for p in doc.paragraphs:
            if not p.text:
                continue
            pPr = p._p.find(qn("w:pPr"))
            if pPr is None:
                continue
            numPr = pPr.find(qn("w:numPr"))
            if numPr is None:
                continue
            ilvl = numPr.find(qn("w:ilvl"))
            if ilvl is None:
                continue
            levels_seen[p.text] = int(ilvl.get(qn("w:val")) or "0")
        # We expect levels 1, 2, 3 to be explicitly stamped (level 0 inherits
        # from the style, no explicit ilvl needed).
        self.assertEqual(levels_seen.get("第二层"), 1)
        self.assertEqual(levels_seen.get("第三层"), 2)
        self.assertEqual(levels_seen.get("第四层"), 3)
        os.unlink(out_path)


class TestRespectTemplateFonts(unittest.TestCase):
    """With a user template, the renderer should not overwrite fonts that
    the template already defines."""

    def _make_template_with_custom_fonts(self, tmp_path):
        """Build a minimal docx where Normal uses Microsoft YaHei for both
        Latin and CJK — the kind of customization a user template would
        carry that the old renderer would silently clobber."""
        from docx import Document as DocxDocument
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn
        d = DocxDocument()
        normal = d.styles["Normal"]
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(_qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(_qn("w:ascii"), "Microsoft YaHei")
        rfonts.set(_qn("w:hAnsi"), "Microsoft YaHei")
        rfonts.set(_qn("w:eastAsia"), "Microsoft YaHei")
        rfonts.set(_qn("w:cs"), "Microsoft YaHei")
        d.save(tmp_path)

    def test_user_template_fonts_preserved(self):
        from md2word.parser import parse
        from md2word.renderer import Renderer
        fd, tpl = tempfile.mkstemp(suffix=".docx"); os.close(fd)
        self._make_template_with_custom_fonts(tpl)

        fd, out = tempfile.mkstemp(suffix=".docx"); os.close(fd)
        ast = parse("这是一段中文正文。\n")
        Renderer(template=tpl).render(ast, out)

        rendered = Document(out)
        p = next(p for p in rendered.paragraphs if p.text)
        # In respect-mode, since the style already declares eastAsia,
        # we should NOT have written a per-run rFonts that overrides it
        # to 宋体. Either there's no run-level rFonts, or it preserves
        # the style's choice.
        for r in p.runs:
            fonts = run_fonts(r)
            if fonts.get("eastAsia"):
                self.assertEqual(fonts["eastAsia"], "Microsoft YaHei",
                                 "user template's CJK font was overwritten")
        os.unlink(tpl); os.unlink(out)

    def test_builtin_template_still_overrides(self):
        """Sanity: with the builtin template (no user template), we still
        apply our default fonts — that's what the existing CJKFontSlots
        test relies on. This is a guard against accidentally flipping the
        default."""
        doc, out_path = render_fixture("cjk_mixed")
        p = next(p for p in doc.paragraphs if p.text)
        for r in p.runs:
            fonts = run_fonts(r)
            self.assertEqual(fonts.get("eastAsia"), "宋体")
        os.unlink(out_path)


class TestInspect(unittest.TestCase):
    """The inspect command reads template metadata without crashing and
    surfaces the kind of warnings users actually need."""

    def test_inspect_builtin_template(self):
        from md2word.inspect import inspect_template, format_report_text
        builtin = os.path.join(os.path.dirname(__file__), "..", "md2word",
                               "templates", "default_zh.docx")
        report = inspect_template(builtin)
        self.assertGreater(report.style_count, 10)
        # All the styles the builtin template was built to provide must match
        roles_matched = {s.role: s.matched for s in report.styles}
        self.assertEqual(roles_matched["一级标题"], "Heading 1")
        self.assertEqual(roles_matched["正文"], "Normal")
        self.assertEqual(roles_matched["代码块"], "Code Block")
        # Page setup parsed
        self.assertIsNotNone(report.page)
        self.assertGreater(report.page.width_mm, 100)
        # Text rendering shouldn't crash
        text = format_report_text(report)
        self.assertIn("模板文件", text)
        self.assertIn("样式匹配", text)

    def test_inspect_warns_on_missing_styles(self):
        """A blank docx (python-docx's built-in default) is missing several
        of our target styles — inspect must say so."""
        from docx import Document as DocxDocument
        from md2word.inspect import inspect_template
        fd, tpl = tempfile.mkstemp(suffix=".docx"); os.close(fd)
        DocxDocument().save(tpl)  # Truly default, no Code Block, no Caption
        report = inspect_template(tpl)
        missing_roles = {s.role for s in report.styles if s.matched is None}
        # Code Block and Caption are not in python-docx's default styles
        self.assertIn("代码块", missing_roles)
        self.assertTrue(any("代码块" in w for w in report.warnings))
        os.unlink(tpl)

    def test_inspect_json_output(self):
        from md2word.inspect import inspect_template, format_report_json
        builtin = os.path.join(os.path.dirname(__file__), "..", "md2word",
                               "templates", "default_zh.docx")
        report = inspect_template(builtin)
        import json
        parsed = json.loads(format_report_json(report))
        self.assertIn("styles", parsed)
        self.assertIn("page", parsed)


class TestCliBackwardCompat(unittest.TestCase):
    """The old `md2word input.md -o out.docx` form must still work even
    though we now have subcommands."""

    def test_implicit_convert_subcommand(self):
        import io, contextlib
        from md2word.cli import main as cli_main
        fd, out = tempfile.mkstemp(suffix=".docx"); os.close(fd)
        md = os.path.join(FIXTURES, "headings.md")
        # No "convert" subcommand on the argv — backwards compat path
        with contextlib.redirect_stdout(io.StringIO()):
            rc = cli_main([md, "-o", out])
        self.assertEqual(rc, 0)
        self.assertTrue(os.path.exists(out))
        os.unlink(out)

    def test_explicit_inspect_subcommand(self):
        import io, contextlib
        from md2word.cli import main as cli_main
        builtin = os.path.join(os.path.dirname(__file__), "..", "md2word",
                               "templates", "default_zh.docx")
        with contextlib.redirect_stdout(io.StringIO()):
            rc = cli_main(["inspect", builtin])
        self.assertEqual(rc, 0)


if __name__ == "__main__":
    unittest.main()
